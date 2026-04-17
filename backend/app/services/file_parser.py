"""
services/file_parser.py — Extract raw text from PDF and DOCX resume files.

Libraries used:
  - pdfplumber for PDF parsing
  - python-docx for DOCX parsing
"""

import os
import pdfplumber
import docx
from fastapi import HTTPException, status


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract all text from a PDF file using pdfplumber.

    Args:
        file_path: Absolute path to the PDF file.

    Returns:
        Concatenated text from all pages.

    Raises:
        HTTPException 422 if the file cannot be read.
    """
    if not os.path.exists(file_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Resume file not found.")

    try:
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Failed to parse PDF: {str(e)}",
        )


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract all text from a DOCX file using python-docx.

    Args:
        file_path: Absolute path to the DOCX file.

    Returns:
        Concatenated text from all paragraphs.

    Raises:
        HTTPException 422 if the file cannot be read.
    """
    if not os.path.exists(file_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Resume file not found.")

    try:
        doc = docx.Document(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        return "\n".join(paragraphs)
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail=f"Failed to parse DOCX: {str(e)}",
        )


def extract_text(file_path: str) -> str:
    """
    Dispatch to the correct parser based on file extension.

    Args:
        file_path: Path to a .pdf or .docx file.

    Returns:
        Extracted raw text.

    Raises:
        HTTPException 415 for unsupported file types.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise HTTPException(
            status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
            detail=f"Unsupported file type '{ext}'. Only PDF and DOCX are accepted.",
        )
